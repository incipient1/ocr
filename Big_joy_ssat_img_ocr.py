
# coding: utf-8

import TencentYoutuyun
import os
import time
import random
import hmac
import hashlib
import binascii
import base64
import requests
import json
from aip import AipOcr # 百度Api
from docx import Document # 用来将结果写入word中的python-docx
import sys

sys.getdefaultencoding()


## 使用腾讯优图


appid = '****'
secret_id = '****'
secret_key = '****'
userid = '****'
end_point = TencentYoutuyun.conf.API_YOUTU_END_POINT
youtu = TencentYoutuyun.YouTu(appid, secret_id, secret_key, userid, end_point)

img_path = 'F:\\NEW-1-1.jpg'

## 腾讯官方接口
ret = youtu.generalocr(img_path, data_type = 0, seq = '')
ret['items'][3]['words'][1]['character'] # 这里显示的是毫无意义的字符

ret['items'][3]['words'][1]['character'].encode('raw_unicode_escape').decode('utf8') # 这里才正常显示了


 ## 使用伟大的giter提供的接口方式[youtu.py](https://github.com/se4/tiny-ocr/blob/0b93a9baadf8dcf563a4c4311a0269efbb3a0846/youtu.py)：

class Youtu(object):

    def __init__(self, app_id, secret_id, secret_key, qq='****'):
        self.app_id = app_id
        self.secret_id = secret_id
        self.secret_key = secret_key
        self.qq = qq

    def cal_sig(self):
        timestamp = int(time.time())
        expired = str(timestamp + 2592000)
        rdm = str(random.randint(0, 999999999))
        plain_text = 'a={appid}&k={secret_id}&e={expired}&t={timestamp}&r={rdm}&u={qq}&f='
        plain_text = plain_text.format(appid=self.app_id,
                                       secret_id=self.secret_id,
                                       timestamp=timestamp,
                                       rdm=rdm, qq=self.qq,
                                       expired=expired)
        bin = hmac.new(self.secret_key.encode(), plain_text.encode(), hashlib.sha1).hexdigest()
        s = binascii.unhexlify(bin)
        s = s + plain_text.encode('ascii')
        signature = base64.b64encode(s).rstrip().decode()
        return signature

    def get_text(self, image_path):
        signature = self.cal_sig()
        headers = {'Host': 'api.youtu.qq.com', 'Content-Type': 'text/json', 'Authorization': signature}
        filepath = os.path.abspath(image_path)
        data = {'app_id': self.app_id, 'image': ''}
        data['image'] = base64.b64encode(open(filepath, 'rb').read()).rstrip().decode('utf-8')
        resp = requests.post('https://api.youtu.qq.com/youtu/ocrapi/generalocr',
                             data=json.dumps(data),
                             headers=headers)
        if 'items' in resp.text:
            return resp.content.decode('utf-8')
        else:
            return '0'


ocr = Youtu(appid, secret_id, secret_key)
resp = ocr.get_text(img_path)
resp = eval(resp)
ret = resp['items'][3]['words'][1]['character']
# 这里很顺畅的生成了汉字
# 但是，我上传的图片中全部都是英文啊！！

## 使用百度云
APP_ID = '****'
API_KEY = '****'
SECRET_KEY = '*****'

client = AipOcr(APP_ID, API_KEY, SECRET_KEY)

# 这里是官方提供的方法
""" 读取图片 """
def get_file_content(filePath):
    with open(filePath, 'rb') as fp:
        return fp.read()

# image = get_file_content('***.jpg')

""" 调用通用文字识别（高精度版） """
# client.basicAccurate(image)

""" 如果有可选参数 """
options = {}
options["detect_direction"] = "true"
options["probability"] = "false"
options["language_type"] = "ENG"

""" 带参数调用通用文字识别（高精度版） """;
# result = client.basicAccurate(image, options)

image = get_file_content('F:\\NEW-1-1.jpg')
result = client.accurate(image, options)

for j in result['words_result']:
    print(j['words'])


# 一眼就能看出来，百度的识别成功率高

### 使用百度云识别整个文件夹中的图片并写到word文档中

path3 = 'F:\\file_2018_05_17\\'
# 读取文件夹中的所有文件，生成文件名list
file = os.listdir(path3)

ok_list = [] # 用来存储有正确返回的
err_list = [] # 用来存储报错的，可用来多次上传百度云再识别
document = Document() # 用来将生成结果写到word文档中去

for fi in file:
    image = get_file_content(path3 + fi)
    result = client.basicAccurate(image, options) # 这里用的是不需要位置信息的接口
    try:
        word_result = result['words_result']
        heading = u'{}'.format(fi)
        document.add_heading(heading,1) # 一级标题

        for i in word_result:
            document.add_paragraph(u'{}'.format(i['words']))

        document.add_page_break() # 两张图片生成的内容之间用分页符
        ok_list.append(fi)

    except:
        err_list.append(fi)
        print(fi, result)

document.save('2018_5_18(5).docx')

# 百度云、腾讯优图的key申请都很简单，但是奇怪的是，再登陆百度云的时候我用chrome浏览器就没有成功，最后转用Edge才可以